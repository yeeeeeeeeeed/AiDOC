"""워드(.docx) 생성 서비스 — 마크다운 텍스트를 docx로 변환"""

import io
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def markdown_to_docx(md_text: str, title: str = "") -> bytes:
    """마크다운 텍스트 → docx 바이트"""
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(10)

    if title:
        doc.add_heading(title, level=0)

    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # 제목
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)

        # 표 (| 로 시작하는 연속 줄)
        elif line.strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            i -= 1
            _add_table(doc, table_lines)

        # 목록
        elif line.strip().startswith("- ") or line.strip().startswith("* "):
            text = re.sub(r'^[\s]*[-*]\s+', '', line)
            doc.add_paragraph(text, style="List Bullet")
        elif re.match(r'^\s*\d+\.\s+', line):
            text = re.sub(r'^\s*\d+\.\s+', '', line)
            doc.add_paragraph(text, style="List Number")

        # 빈 줄
        elif line.strip() == "":
            pass

        # 일반 텍스트
        else:
            p = doc.add_paragraph()
            _add_formatted_text(p, line)

        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_table(doc: Document, lines: list[str]):
    """마크다운 표 → docx 테이블"""
    rows_data = []
    for line in lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        # 구분선 (---) 건너뛰기
        if all(re.match(r'^[-:]+$', c) for c in cells):
            continue
        rows_data.append(cells)

    if not rows_data:
        return

    cols = max(len(r) for r in rows_data)
    table = doc.add_table(rows=len(rows_data), cols=cols)
    table.style = "Table Grid"

    for r, row_data in enumerate(rows_data):
        for c, val in enumerate(row_data):
            if c < cols:
                table.cell(r, c).text = val

    # 첫 행 볼드
    if rows_data:
        for cell in table.rows[0].cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True


def _add_formatted_text(paragraph, text: str):
    """볼드(**), 이탤릭(*) 처리"""
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)
